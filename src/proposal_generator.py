# src/proposal_generator.py

import os
from docx import Document
from src.agentic_orchestrator import agentic_rfp_answer, FALLBACK_RESPONSE


def generate_proposal_from_sections(
    sections,
    output_path,
    provider="ollama",
    resume=True
):
    """
    Generates a structured RFP proposal Word document from parsed sections.

    resume=True: If output file exists, skips already-processed headings.
                 Note: only skips if heading text AND a following paragraph exist.
                 Incomplete sections will be regenerated.
    """

    # -------------------------------
    # Resume handling
    # -------------------------------

    if resume and os.path.exists(output_path):
        print("Resuming existing document...")
        doc = Document(output_path)

        # Build set of headings that have content after them (i.e. fully written)
        existing_headings = set()
        paragraphs = doc.paragraphs

        for i, para in enumerate(paragraphs):
            if para.style.name.startswith("Heading"):
                # Check if the next paragraph exists and has content
                if i + 1 < len(paragraphs) and paragraphs[i + 1].text.strip():
                    existing_headings.add(para.text.strip())

        print(f"   → Found {len(existing_headings)} completed sections to skip.\n")

    else:
        doc = Document()
        existing_headings = set()

    total_sections = len(sections)

    print("====================================")
    print("Starting Proposal Generation")
    print("====================================")
    print(f"Total sections detected: {total_sections}\n")

    completed = 0
    skipped = 0
    errors = 0

    try:
        for idx, (heading, bullets) in enumerate(sections, 1):

            print(f"[{idx}/{total_sections}] Section: {heading}")

            if heading.strip() in existing_headings:
                print("   → Already completed. Skipping.\n")
                skipped += 1
                continue

            combined_requirements = "\n".join(
                f"- {b.strip()}" for b in bullets if b.strip()
            )

            if not combined_requirements.strip():
                print("   → No bullet content found. Skipping.\n")
                skipped += 1
                continue

            query = f"""
Respond to the following section requirements in a structured enterprise manner.

Section:
{heading}

Requirements:
{combined_requirements}
"""

            doc.add_heading(heading, level=1)

            print("   → Generating response...")

            try:
                response = agentic_rfp_answer(
                    query,
                    mode="proposal",
                    provider=provider
                )
            except Exception as e:
                print(f"   ⚠ Generation failed for section '{heading}': {e}")
                response = FALLBACK_RESPONSE
                errors += 1

            if not response:
                response = FALLBACK_RESPONSE

            doc.add_paragraph(response)
            doc.save(output_path)  # Save after every section for crash safety

            print("   → Section saved.\n")
            completed += 1

    except KeyboardInterrupt:
        print("\n⚠ Interrupted by user. Saving progress...")
        doc.save(output_path)
        print(f"Progress saved to: {output_path}")
        raise

    except Exception as e:
        print(f"\n⚠ Unexpected error: {e}. Saving progress...")
        doc.save(output_path)
        raise

    print("====================================")
    print("Proposal generation complete.")
    print(f"Completed : {completed} | Skipped : {skipped} | Errors : {errors}")
    print(f"Output    : {output_path}")
    print("====================================")