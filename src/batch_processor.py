# src/batch_processor.py

import os
import pandas as pd
from docx import Document
from src.agentic_orchestrator import agentic_rfp_answer


def _is_section_header(row, df):
    """
    Detects whether a row is a section header rather than a feature row.

    Criteria:
    - Text ends with ":" OR
    - Compliance and Remarks columns are both empty AND text is short (< 60 chars)
    """
    text = str(row.get("Feature / Capability", "")).strip()

    if text.endswith(":"):
        return True

    compliance_val = str(row.get("Compliance (Yes/No/Partial)", "")).strip()
    remarks_val = str(row.get("Remarks / Notes  (For Partial put remarks of capability)", "")).strip()

    if compliance_val in ["", "nan"] and remarks_val in ["", "nan"] and len(text) < 60:
        return True

    return False


def process_compliance_sheet(
    input_file,
    output_file,
    word_output,
    provider="ollama"
):
    """
    Processes Excel compliance sheet row by row.
    Fills Compliance (YES/NO/PARTIAL) and Remarks columns.
    Also generates a Word summary document.

    Errors on individual rows are caught and logged — processing continues.
    """

    print("\nProcessing compliance sheet...\n")

    excel = pd.ExcelFile(input_file)
    writer = pd.ExcelWriter(output_file, engine="openpyxl")
    word_doc = Document()

    for sheet_name in excel.sheet_names:

        print(f"Processing sheet: {sheet_name}")
        df = excel.parse(sheet_name)
        df.columns = df.columns.astype(str).str.strip()

        if "Feature / Capability" not in df.columns:
            print("   → Skipping sheet (no 'Feature / Capability' column found)")
            df.to_excel(writer, sheet_name=sheet_name, index=False)
            continue

        # Ensure output columns exist
        if "Compliance (Yes/No/Partial)" not in df.columns:
            df["Compliance (Yes/No/Partial)"] = ""

        if "Remarks / Notes  (For Partial put remarks of capability)" not in df.columns:
            df["Remarks / Notes  (For Partial put remarks of capability)"] = ""

        current_section = ""
        processed = 0
        skipped = 0
        errors = 0

        for idx, row in df.iterrows():

            text = str(row["Feature / Capability"]).strip()

            if not text or text.lower() == "nan":
                skipped += 1
                continue

            # Detect section headers
            if _is_section_header(row, df):
                current_section = text.rstrip(":")
                print(f"   → Section: {current_section}")
                continue

            enriched_query = f"{current_section} {text}".strip() if current_section else text

            print(f"   → Row {idx + 2}: {text[:60]}")

            # Per-row error handling — one bad row won't abort the sheet
            try:
                response = agentic_rfp_answer(
                    enriched_query,
                    mode="compliance",
                    provider=provider
                )
            except Exception as e:
                print(f"   ⚠ Row {idx + 2} failed: {e}")
                errors += 1
                continue

            if not response:
                skipped += 1
                continue

            lines = response.split("\n")
            compliance = lines[0].strip().upper()

            if compliance not in ["YES", "NO", "PARTIAL"]:
                compliance = "PARTIAL"

            remarks = "\n".join(lines[1:]).strip()

            df.at[idx, "Compliance (Yes/No/Partial)"] = compliance
            df.at[idx, "Remarks / Notes  (For Partial put remarks of capability)"] = remarks

            # Word summary
            word_doc.add_heading(text, level=2)
            word_doc.add_paragraph(response)

            processed += 1

        df.to_excel(writer, sheet_name=sheet_name, index=False)

        print(f"\n   Sheet '{sheet_name}' complete — "
              f"Processed: {processed} | Skipped: {skipped} | Errors: {errors}\n")

    writer.close()
    word_doc.save(word_output)

    print("\nCompliance processing complete.")
    print(f"Excel output : {output_file}")
    print(f"Word summary : {word_output}")